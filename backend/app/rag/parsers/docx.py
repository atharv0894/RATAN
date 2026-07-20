from .base import BaseParser
from .models import ParsedDocument, ParsedPage

class DOCXParser(BaseParser):
    def parse(self, file_path: str, filename: str, **kwargs) -> ParsedDocument:
        metadata = self.extract_metadata(file_path)
        
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            # Fallback if python-docx not installed
            text = f"[DOCX Content of {filename}. Please install python-docx for full extraction]"
            
        metadata["page_count"] = 1
        metadata["mime_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        page = ParsedPage(page_number=1, text=text)
        
        return ParsedDocument(
            filename=filename,
            metadata=metadata,
            pages=[page],
            text=text
        )
